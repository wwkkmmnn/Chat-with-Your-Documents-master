from __future__ import annotations

from pathlib import Path

from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document

try:
    import docx  # type: ignore
except ImportError:  # pragma: no cover - optional dependency
    docx = None

# 根据文件路径的后缀名选择合适的解析方法，支持 PDF、TXT 和 DOCX 格式，返回一个 Document 对象列表，每个对象包含文本内容和元信息（如来源文件名和页码），供后续处理使用
def parse_file(file_path: Path) -> list[Document]:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        documents = PyPDFLoader(str(file_path)).load()
        for document in documents:
            document.metadata.setdefault("source", file_path.name)
        return documents

    if suffix == ".txt":
        text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
        return [Document(page_content=text, metadata={"source": file_path.name, "page": 0})]

    if suffix == ".docx":
        if docx is None:
            raise ValueError("DOCX parsing requires python-docx.")

        document = docx.Document(str(file_path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        content = "\n".join(paragraphs).strip()
        return [Document(page_content=content, metadata={"source": file_path.name, "page": 0})]

    raise ValueError("Unsupported file type.")
